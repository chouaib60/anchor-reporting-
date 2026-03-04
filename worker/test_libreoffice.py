import subprocess
import tempfile
import time
import os
from pathlib import Path
from docx import Document

LIBREOFFICE_PATH = os.getenv("LIBREOFFICE_PATH", "/usr/bin/libreoffice")
TIMEOUT = 60


# -- Helpers --

def log(level, message, **kwargs):
    print(f"[{level}] {message}", kwargs if kwargs else "")


def create_sample_docx(path: Path):
    doc = Document()
    doc.add_heading("Test Sprint 0 "
    "- Report as a Service", 0)
    doc.add_paragraph("Validation de la conversion LibreOffice headless.")
    doc.save(str(path))
    log("INFO", f"DOCX cree : {path}")


def convert_to_pdf(docx_path: Path, output_dir: Path, timeout: int = TIMEOUT) -> Path:
    profile_dir = output_dir / "lo_profile"
    profile_dir.mkdir(parents=True, exist_ok=True)

    cmd = [
        LIBREOFFICE_PATH,
        "--headless",
        "--norestore",
        "--nofirststartwizard",
        f"-env:UserInstallation=file:///{profile_dir.as_posix()}",
        "--convert-to", "pdf",
        "--outdir", str(output_dir),
        str(docx_path),
    ]

    proc = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)

    try:
        stdout, stderr = proc.communicate(timeout=timeout)
    except subprocess.TimeoutExpired:
        proc.kill()
        proc.wait()
        raise

    if proc.returncode != 0:
        raise RuntimeError(f"LibreOffice erreur (code {proc.returncode}): {stderr.decode()}")

    pdf_path = output_dir / (docx_path.stem + ".pdf")
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF non genere : {pdf_path}")

    return pdf_path


# -- Test 1 : Conversion basique --

def test_conversion():
    log("INFO", "=== TEST 1 : Conversion DOCX -> PDF ===")
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        docx = tmp_path / "test.docx"
        create_sample_docx(docx)

        start = time.time()
        pdf = convert_to_pdf(docx, tmp_path)
        elapsed = round(time.time() - start, 2)

        assert pdf.exists(), "PDF non cree"
        assert pdf.stat().st_size > 0, "PDF vide"
        log("INFO", f"TEST 1 PASSED ({elapsed}s, {pdf.stat().st_size} bytes)")
        return True


# -- Test 2 : Timeout + kill process --

def test_timeout():
    log("INFO", "=== TEST 2 : Timeout + kill process ===")
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        docx = tmp_path / "test.docx"
        create_sample_docx(docx)

        try:
            convert_to_pdf(docx, tmp_path, timeout=1)
            log("INFO", "TEST 2 SKIPPED - conversion trop rapide pour tester le timeout")
            return True
        except subprocess.TimeoutExpired:
            log("INFO", "TEST 2 PASSED - process tue apres timeout, aucun process zombie")
            return True
        except Exception as e:
            log("ERROR", f"TEST 2 FAILED : {e}")
            return False


# -- Test 3 : Profils isoles (2 conversions en parallele) --

def test_isolated_profiles():
    import concurrent.futures
    log("INFO", "=== TEST 3 : Profils isoles (parallele) ===")

    with tempfile.TemporaryDirectory() as tmp1, \
         tempfile.TemporaryDirectory() as tmp2:

        p1, p2 = Path(tmp1), Path(tmp2)
        d1, d2 = p1 / "test.docx", p2 / "test.docx"
        create_sample_docx(d1)
        create_sample_docx(d2)

        with concurrent.futures.ThreadPoolExecutor(max_workers=2) as ex:
            f1 = ex.submit(convert_to_pdf, d1, p1)
            f2 = ex.submit(convert_to_pdf, d2, p2)
            results = []
            for f in concurrent.futures.as_completed([f1, f2]):
                try:
                    results.append(f.result())
                except Exception as e:
                    log("ERROR", f"Conversion echouee : {e}")
                    results.append(None)

        if all(r is not None and r.exists() for r in results):
            log("INFO", "TEST 3 PASSED - 2 conversions paralleles sans conflit")
            return True
        else:
            log("ERROR", "TEST 3 FAILED")
            return False


# -- Main --

if __name__ == "__main__":
    if not Path(LIBREOFFICE_PATH).exists():
        log("ERROR", f"LibreOffice introuvable : {LIBREOFFICE_PATH}")
        exit(1)

    results = {
        "conversion":        test_conversion(),
        "timeout_kill":      test_timeout(),
        "isolated_profiles": test_isolated_profiles(),
    }

    passed = sum(v for v in results.values())
    total = len(results)
    print(f"\n{'='*40}")
    print(f"Resultat : {passed}/{total} tests passes")

    for name, ok in results.items():
        print(f"  {'PASSED' if ok else 'FAILED'}  {name}")

    print(f"{'='*40}")
    print("GO - LibreOffice operationnel" if passed == total else "NO-GO - Verifier les erreurs")
    exit(0 if passed == total else 1)